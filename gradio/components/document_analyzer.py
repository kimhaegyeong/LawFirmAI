#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Analyzer Component for LawFirmAI
PDF/DOCX parsing and contract analysis functionality
"""

import os
import logging
from pathlib import Path
from typing import Dict, Any, List, Optional
import re

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    logging.warning("PyPDF2 not installed. PDF parsing will not be available.")

try:
    from docx import Document
except ImportError:
    Document = None
    logging.warning("python-docx not installed. DOCX parsing will not be available.")

logger = logging.getLogger(__name__)

class DocumentAnalyzer:
    """문서 분석 클래스 - PDF/DOCX 파싱 및 계약서 분석"""
    
    def __init__(self, rag_service=None):
        """
        DocumentAnalyzer 초기화
        
        Args:
            rag_service: RAG 서비스 인스턴스 (선택사항)
        """
        self.rag_service = rag_service
        
        # 위험 키워드 정의
        self.risk_keywords = {
            "high": [
                "손해배상", "위약금", "해지", "면책", "책임", "배상", "손실",
                "위험", "부담", "책임제한", "면책조항", "해지권", "위약금액"
            ],
            "medium": [
                "계약", "조건", "기간", "연장", "갱신", "변경", "수정",
                "통지", "고지", "이행", "불이행", "지연", "연체"
            ],
            "low": [
                "당사자", "목적", "기간", "대가", "지급", "수령", "인도",
                "인수", "보관", "관리", "운영", "사용", "이용"
            ]
        }
        
        # 법률 용어 패턴
        self.legal_patterns = {
            "contract_terms": [
                r"제\d+조", r"제\d+항", r"제\d+호",  # 조항 번호
                r"계약기간", r"계약조건", r"계약내용",  # 계약 관련
                r"당사자", r"갑", r"을", r"병", r"정"  # 당사자 표시
            ],
            "obligations": [
                r"의무", r"책임", r"이행", r"준수", r"보장",
                r"인도", r"인수", r"지급", r"수령", r"보관"
            ],
            "penalties": [
                r"위약금", r"손해배상", r"배상", r"손실", r"보상",
                r"과태료", r"벌금", r"처벌", r"제재"
            ]
        }
    
    def parse_pdf(self, file_path: str) -> str:
        """
        PDF 파일에서 텍스트 추출
        
        Args:
            file_path: PDF 파일 경로
            
        Returns:
            추출된 텍스트
        """
        if PyPDF2 is None:
            raise ImportError("PyPDF2가 설치되지 않았습니다. pip install PyPDF2로 설치하세요.")
        
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += f"\n--- 페이지 {page_num + 1} ---\n"
                            text += page_text
                    except Exception as e:
                        logger.warning(f"페이지 {page_num + 1} 텍스트 추출 실패: {e}")
                        continue
                
                return text.strip()
                
        except Exception as e:
            logger.error(f"PDF 파싱 오류: {e}")
            raise ValueError(f"PDF 파일을 읽을 수 없습니다: {str(e)}")
    
    def parse_docx(self, file_path: str) -> str:
        """
        DOCX 파일에서 텍스트 추출
        
        Args:
            file_path: DOCX 파일 경로
            
        Returns:
            추출된 텍스트
        """
        if Document is None:
            raise ImportError("python-docx가 설치되지 않았습니다. pip install python-docx로 설치하세요.")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # 표 내용도 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX 파싱 오류: {e}")
            raise ValueError(f"DOCX 파일을 읽을 수 없습니다: {str(e)}")
    
    def parse_document(self, file_path: str) -> str:
        """
        파일 확장자에 따라 적절한 파서 선택
        
        Args:
            file_path: 파일 경로
            
        Returns:
            추출된 텍스트
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")
        
        if file_path.suffix.lower() == '.pdf':
            return self.parse_pdf(str(file_path))
        elif file_path.suffix.lower() == '.docx':
            return self.parse_docx(str(file_path))
        else:
            raise ValueError(f"지원하지 않는 파일 형식입니다: {file_path.suffix}")
    
    def analyze_contract(self, text: str) -> Dict[str, Any]:
        """
        계약서 텍스트 분석
        
        Args:
            text: 분석할 텍스트
            
        Returns:
            분석 결과 딕셔너리
        """
        try:
            # 기본 정보 추출
            summary = self._generate_summary(text)
            
            # 조항 추출
            clauses = self._extract_clauses(text)
            
            # 위험 요소 평가
            risks = self._assess_risks(text, clauses)
            
            # 개선 제안 생성
            recommendations = self._generate_recommendations(risks, text)
            
            # 전체 위험도 점수 계산
            risk_score = self._calculate_risk_score(risks)
            
            return {
                "summary": summary,
                "clauses": clauses,
                "risks": risks,
                "recommendations": recommendations,
                "risk_score": risk_score,
                "analysis_metadata": {
                    "text_length": len(text),
                    "clause_count": len(clauses),
                    "risk_count": len(risks),
                    "high_risk_count": len([r for r in risks if r["risk_level"] == "high"])
                }
            }
            
        except Exception as e:
            logger.error(f"계약서 분석 오류: {e}")
            raise ValueError(f"문서 분석 중 오류가 발생했습니다: {str(e)}")
    
    def _generate_summary(self, text: str) -> str:
        """문서 요약 생성"""
        lines = text.split('\n')
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        
        # 문서 길이 정보
        char_count = len(text)
        word_count = len(text.split())
        line_count = len(non_empty_lines)
        
        # 주요 키워드 추출
        keywords = self._extract_keywords(text)
        
        summary = f"""**문서 기본 정보**
- 문서 길이: {char_count:,}자
- 단어 수: {word_count:,}개
- 유효 라인 수: {line_count:,}개

**주요 키워드**
{', '.join(keywords[:10]) if keywords else '키워드를 찾을 수 없습니다.'}

**분석 완료**: {len(non_empty_lines)}개 섹션에서 정보를 추출했습니다."""
        
        return summary
    
    def _extract_clauses(self, text: str) -> List[Dict]:
        """중요한 조항 추출"""
        clauses = []
        
        # 조항 번호 패턴으로 조항 찾기
        clause_pattern = r'(제\s*\d+\s*조[^\n]*)'
        matches = re.finditer(clause_pattern, text, re.IGNORECASE)
        
        for i, match in enumerate(matches, 1):
            clause_text = match.group(1).strip()
            if len(clause_text) > 10:  # 너무 짧은 조항 제외
                clauses.append({
                    "id": i,
                    "title": f"제 {i}조",
                    "content": clause_text[:200] + "..." if len(clause_text) > 200 else clause_text,
                    "type": "numbered_clause",
                    "importance": self._assess_clause_importance(clause_text)
                })
        
        # RAG 서비스가 있으면 추가 분석
        if self.rag_service and clauses:
            try:
                # 상위 3개 조항에 대해 RAG 분석
                for clause in clauses[:3]:
                    query = f"다음 계약 조항의 법적 의미와 주의사항을 설명해주세요: {clause['content'][:100]}"
                    rag_result = self.rag_service.process_query(query, top_k=3)
                    
                    if rag_result and rag_result.get("sources"):
                        clause["rag_analysis"] = rag_result["sources"][0].get("content", "")[:200]
            except Exception as e:
                logger.warning(f"RAG 분석 실패: {e}")
        
        return clauses
    
    def _assess_clause_importance(self, clause_text: str) -> str:
        """조항 중요도 평가"""
        high_importance_keywords = ["손해배상", "위약금", "해지", "면책", "책임"]
        medium_importance_keywords = ["계약", "조건", "기간", "이행", "통지"]
        
        clause_lower = clause_text.lower()
        
        if any(keyword in clause_lower for keyword in high_importance_keywords):
            return "high"
        elif any(keyword in clause_lower for keyword in medium_importance_keywords):
            return "medium"
        else:
            return "low"
    
    def _assess_risks(self, text: str, clauses: List[Dict]) -> List[Dict]:
        """위험 요소 평가"""
        risks = []
        
        # 텍스트 전체에서 위험 키워드 검색
        text_lower = text.lower()
        
        for risk_level, keywords in self.risk_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    # 키워드 주변 컨텍스트 추출
                    context = self._extract_context(text, keyword, 100)
                    
                    risks.append({
                        "keyword": keyword,
                        "risk_level": risk_level,
                        "context": context,
                        "reason": self._get_risk_reason(keyword, risk_level),
                        "recommendation": self._get_risk_recommendation(keyword, risk_level)
                    })
        
        # 중복 제거 및 정렬
        unique_risks = []
        seen_keywords = set()
        
        for risk in risks:
            if risk["keyword"] not in seen_keywords:
                unique_risks.append(risk)
                seen_keywords.add(risk["keyword"])
        
        # 위험도 순으로 정렬
        risk_order = {"high": 3, "medium": 2, "low": 1}
        unique_risks.sort(key=lambda x: risk_order.get(x["risk_level"], 0), reverse=True)
        
        return unique_risks[:10]  # 상위 10개만 반환
    
    def _extract_context(self, text: str, keyword: str, context_length: int = 100) -> str:
        """키워드 주변 컨텍스트 추출"""
        try:
            index = text.lower().find(keyword.lower())
            if index == -1:
                return ""
            
            start = max(0, index - context_length)
            end = min(len(text), index + len(keyword) + context_length)
            
            context = text[start:end]
            return context.strip()
        except Exception:
            return ""
    
    def _get_risk_reason(self, keyword: str, risk_level: str) -> str:
        """위험 요소 이유 설명"""
        reasons = {
            "high": f"'{keyword}' 키워드가 포함되어 있어 높은 법적 위험이 있습니다.",
            "medium": f"'{keyword}' 관련 조항이 있어 주의가 필요합니다.",
            "low": f"'{keyword}' 관련 내용이 포함되어 있습니다."
        }
        return reasons.get(risk_level, "위험 요소가 발견되었습니다.")
    
    def _get_risk_recommendation(self, keyword: str, risk_level: str) -> str:
        """위험 요소별 개선 제안"""
        recommendations = {
            "high": f"'{keyword}' 조항에 대해 법무팀 검토를 권장합니다.",
            "medium": f"'{keyword}' 관련 내용을 명확히 정의하는 것이 좋습니다.",
            "low": f"'{keyword}' 관련 조항을 확인해보세요."
        }
        return recommendations.get(risk_level, "전문가 검토를 권장합니다.")
    
    def _generate_recommendations(self, risks: List[Dict], text: str) -> List[str]:
        """개선 제안 생성"""
        recommendations = []
        
        # 위험도별 개선 제안
        high_risks = [r for r in risks if r["risk_level"] == "high"]
        medium_risks = [r for r in risks if r["risk_level"] == "medium"]
        
        if high_risks:
            recommendations.append(f"⚠️ **높은 위험도 조항 {len(high_risks)}개 발견**")
            recommendations.append("법무팀의 전문 검토가 필요합니다.")
            
            for risk in high_risks[:3]:  # 상위 3개만
                recommendations.append(f"- {risk['keyword']}: {risk['recommendation']}")
        
        if medium_risks:
            recommendations.append(f"⚡ **중간 위험도 조항 {len(medium_risks)}개 발견**")
            recommendations.append("조항 내용을 명확히 정의하는 것을 권장합니다.")
        
        # 일반적인 개선 제안
        if not high_risks and not medium_risks:
            recommendations.append("✅ 특별한 위험 요소가 발견되지 않았습니다.")
        
        # 문서 품질 개선 제안
        recommendations.extend([
            "",
            "📋 **일반적인 개선 제안**",
            "• 계약 당사자 정보를 명확히 기재하세요",
            "• 계약 목적과 범위를 구체적으로 명시하세요",
            "• 이행 기간과 방법을 상세히 기술하세요",
            "• 분쟁 해결 방법을 명시하세요"
        ])
        
        return recommendations
    
    def _calculate_risk_score(self, risks: List[Dict]) -> int:
        """전체 위험도 점수 계산 (0-100)"""
        if not risks:
            return 0
        
        # 위험도별 가중치
        weights = {"high": 3, "medium": 2, "low": 1}
        
        total_weight = sum(weights.get(risk["risk_level"], 1) for risk in risks)
        max_possible_weight = len(risks) * 3  # 모든 위험이 high일 때
        
        if max_possible_weight == 0:
            return 0
        
        # 0-100 점수로 변환
        score = int((total_weight / max_possible_weight) * 100)
        return min(100, score)
    
    def _extract_keywords(self, text: str) -> List[str]:
        """주요 키워드 추출"""
        # 간단한 키워드 추출 (실제로는 더 정교한 방법 사용 가능)
        words = re.findall(r'\b[가-힣]{2,}\b', text)
        
        # 빈도 계산
        word_freq = {}
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # 상위 키워드 반환
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:20] if freq > 1]
